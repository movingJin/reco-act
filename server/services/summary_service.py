from typing import List, Dict, Any, Optional
from database import SessionLocal, Meeting as DBMeeting, Paragraph as DBParagraph, NextStep as DBNextStep
from graph.summary_workflow import SummaryNode, create_meeting_notes
from models.summary import SummaryResponse, Paragraph


def get_db():
    """Get database session."""
    return SessionLocal()


def db_summary_to_model(db_meeting: DBMeeting, db_paragraphs: List[DBParagraph], db_next_steps: List[DBNextStep]) -> SummaryResponse:
    """Convert database models to Pydantic SummaryResponse model."""
    # Sort paragraphs by start time (ascending)
    sorted_paragraphs = sorted(db_paragraphs, key=lambda p: p.start)
    
    paragraphs = [
        Paragraph(
            id=p.id,
            subject=p.subject,
            summary=p.summary,
            start=p.start,
            end=p.end
        )
        for p in sorted_paragraphs
    ]
    
    # order_index(삽입 순서) 기준으로 정렬해 반환한다. 과거 데이터로 값이 없으면 맨 뒤로.
    sorted_next_steps = sorted(
        db_next_steps,
        key=lambda s: s.order_index if s.order_index is not None else len(db_next_steps)
    )
    next_steps = [step.todo for step in sorted_next_steps]

    return SummaryResponse(
        meeting_id=db_meeting.id,
        paragraphs=paragraphs,
        next_steps=next_steps,
        subject=db_meeting.subject,
        meeting_notes=db_meeting.meeting_notes
    )


def load_summary(meeting_id: str) -> Optional[SummaryResponse]:
    """Load a summary from database."""
    db = get_db()
    try:
        db_meeting = db.query(DBMeeting).filter(DBMeeting.id == meeting_id).first()
        if not db_meeting:
            return None
        
        db_paragraphs = db.query(DBParagraph).filter(DBParagraph.meeting_id == meeting_id).all()
        db_next_steps = db.query(DBNextStep).filter(DBNextStep.meeting_id == meeting_id).all()
        return db_summary_to_model(db_meeting, db_paragraphs, db_next_steps)
    except Exception as e:
        print(f"Error loading meeting {meeting_id}: {e}")
        return None
    finally:
        db.close()


def create_summary(meeting_id: str) -> Optional[SummaryResponse]:
    """Create a new summary in database."""
    db = get_db()
    try:
        summary_node = SummaryNode()
        summaries = summary_node.run(meeting_id)
        
        # Delete existing paragraphs and next steps for this meeting
        db.query(DBParagraph).filter(DBParagraph.meeting_id == meeting_id).delete()
        db.query(DBNextStep).filter(DBNextStep.meeting_id == meeting_id).delete()
        
        # Add paragraphs to database
        db_paragraphs_list = []
        for p in summaries.paragraphs:
            db_paragraph = DBParagraph(
                meeting_id=meeting_id,
                subject=p["subject"],
                start=p["start"],
                end=p["end"],
                summary=p["summary"]
            )
            db.add(db_paragraph)
            db_paragraphs_list.append(db_paragraph)
        
        # Add next steps to database (order_index로 순서 보존)
        for idx, step in enumerate(summaries.next_steps):
            db.add(DBNextStep(
                meeting_id=meeting_id,
                todo=step,
                order_index=idx
            ))
        
        # Update meeting subject and generated meeting notes
        db_meeting = db.query(DBMeeting).filter(DBMeeting.id == meeting_id).first()
        db_meeting.subject = summaries.subject
        db_meeting.meeting_notes = summaries.meeting_notes
        db.commit()
        
        # Return updated summary (this will apply sorting via db_summary_to_model)
        db_paragraphs = db.query(DBParagraph).filter(DBParagraph.meeting_id == meeting_id).all()
        db_next_steps = db.query(DBNextStep).filter(DBNextStep.meeting_id == meeting_id).all()
        return db_summary_to_model(db_meeting, db_paragraphs, db_next_steps)
    except Exception as e:
        print(f"Error creating meeting: {e}")
        db.rollback()
        raise
    finally:
        db.close()


def update_paragraph(meeting_id: str, paragraph_id: int, paragraph_data: Dict[str, Any]) -> Optional[SummaryResponse]:
    """Update a specific paragraph of a meeting in database."""
    db = get_db()
    try:
        # Get the paragraph to update
        db_paragraph = db.query(DBParagraph).filter(
            DBParagraph.id == paragraph_id,
            DBParagraph.meeting_id == meeting_id
        ).first()
        
        if not db_paragraph:
            return None
        
        # Update paragraph fields
        if "subject" in paragraph_data:
            db_paragraph.subject = paragraph_data["subject"]
        if "summary" in paragraph_data:
            db_paragraph.summary = paragraph_data["summary"]
        if "start" in paragraph_data:
            db_paragraph.start = paragraph_data["start"]
        if "end" in paragraph_data:
            db_paragraph.end = paragraph_data["end"]
        
        db.commit()
        
        # Return updated summary
        db_meeting = db.query(DBMeeting).filter(DBMeeting.id == meeting_id).first()
        db_paragraphs = db.query(DBParagraph).filter(DBParagraph.meeting_id == meeting_id).all()
        db_next_steps = db.query(DBNextStep).filter(DBNextStep.meeting_id == meeting_id).all()
        
        return db_summary_to_model(db_meeting, db_paragraphs, db_next_steps)
    except Exception as e:
        print(f"Error updating paragraph: {e}")
        db.rollback()
        return None
    finally:
        db.close()


def update_next_steps(meeting_id: str, next_steps: List[str]) -> Optional[SummaryResponse]:
    """Update the next steps of a meeting in database."""
    db = get_db()
    try:
        db_meeting = db.query(DBMeeting).filter(DBMeeting.id == meeting_id).first()
        if not db_meeting:
            return None
        
        # Delete existing next steps for this meeting
        db.query(DBNextStep).filter(DBNextStep.meeting_id == meeting_id).delete()
        
        # Add new next steps (order_index로 순서 보존)
        for idx, step in enumerate(next_steps):
            db.add(DBNextStep(
                meeting_id=meeting_id,
                todo=step,
                order_index=idx
            ))
        
        db.commit()
        
        # Return updated summary
        db_paragraphs = db.query(DBParagraph).filter(DBParagraph.meeting_id == meeting_id).all()
        db_next_steps_updated = db.query(DBNextStep).filter(DBNextStep.meeting_id == meeting_id).all()
        
        return db_summary_to_model(db_meeting, db_paragraphs, db_next_steps_updated)
    except Exception as e:
        print(f"Error updating next steps: {e}")
        db.rollback()
        return None
    finally:
        db.close()


def update_meeting_notes(meeting_id: str, meeting_notes: str) -> Optional[SummaryResponse]:
    """회의록(meeting_notes) 본문을 DB에 저장합니다(미리보기에서 편집한 내용 반영)."""
    db = get_db()
    try:
        db_meeting = db.query(DBMeeting).filter(DBMeeting.id == meeting_id).first()
        if not db_meeting:
            return None

        db_meeting.meeting_notes = meeting_notes
        db.commit()

        db_paragraphs = db.query(DBParagraph).filter(DBParagraph.meeting_id == meeting_id).all()
        db_next_steps = db.query(DBNextStep).filter(DBNextStep.meeting_id == meeting_id).all()
        return db_summary_to_model(db_meeting, db_paragraphs, db_next_steps)
    except Exception as e:
        print(f"Error updating meeting notes: {e}")
        db.rollback()
        return None
    finally:
        db.close()


def regenerate_meeting_notes(meeting_id: str) -> Optional[SummaryResponse]:
    """기존 요약(단락/주제/다음 할 일)은 그대로 두고 회의록 본문만 다시 생성한다."""
    summary = load_summary(meeting_id)
    if not summary:
        return None

    # 회의록 노드에 넘길 state 구성 (DB에 저장된 기존 요약을 그대로 사용)
    state = {
        "meeting_id": meeting_id,
        "subject": summary.subject,
        "next_steps": summary.next_steps,
        "paragraphs": [
            {"subject": p.subject, "start": p.start, "end": p.end, "summary": p.summary}
            for p in summary.paragraphs
        ],
        "meeting_notes": "",
        "current_step": "regenerate_meeting_notes",
    }

    try:
        result_state = create_meeting_notes(state)
    except Exception as e:
        print(f"Error regenerating meeting notes: {e}")
        return None

    return update_meeting_notes(meeting_id, result_state.get("meeting_notes") or "")


def generate_meeting_notes_docx(summary: SummaryResponse, meeting=None) -> bytes:
    """회의록(meeting_notes)을 Word(.docx) 문서 바이트로 생성합니다."""
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt

    document = Document()

    title = (meeting.title if meeting and meeting.title else None) or "회의록"
    document.add_heading(title, level=0)

    notes = (summary.meeting_notes or "").strip()
    if not notes:
        document.add_paragraph("(작성된 회의록이 없습니다)")
    else:
        for line in notes.split("\n"):
            stripped = line.strip()
            # "1. ", "2. " 처럼 숫자로 시작하는 최상위 섹션 제목은 굵게 표시한다.
            is_section = len(stripped) > 1 and stripped[0].isdigit() and stripped[1:3] in (". ", ".")
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            if is_section:
                run.bold = True
                run.font.size = Pt(13)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def format_time(ms: int) -> str:
    """Format milliseconds to HH:MM:SS format."""
    total_seconds = ms // 1000
    hours = total_seconds // 3600
    minutes = (total_seconds % 3600) // 60
    seconds = total_seconds % 60
    return f"{hours}:{str(minutes).zfill(2)}:{str(seconds).zfill(2)}"


def generate_summary_text(summary: SummaryResponse, meeting=None) -> str:
    """Generate formatted summary text for download."""
    summary_text = ""

    # Add header
    if meeting is not None:
        summary_text += f"회의록: {meeting.title}\n"
        summary_text += f"생성일: {meeting.created_at}\n"
        summary_text += f"참석자: {', '.join(meeting.participants)}\n"
        summary_text += "=" * 60 + "\n\n"
    else:
        summary_text += "=" * 60 + "\n"
        summary_text += "회의 요약\n"
        summary_text += "=" * 60 + "\n\n"

    # Add subject section
    summary_text += "[요약]\n"
    summary_text += "-" * 60 + "\n"
    summary_text += summary.subject or "(요약 없음)\n"
    summary_text += "\n\n"
    
    # Add next steps section
    summary_text += "[다음 할 일]\n"
    summary_text += "-" * 60 + "\n"
    if summary.next_steps and len(summary.next_steps) > 0:
        for idx, step in enumerate(summary.next_steps, 1):
            summary_text += f"{idx}. {step}\n"
    else:
        summary_text += "(할 일 없음)\n"
    summary_text += "\n\n"
    
    # Add paragraphs section
    summary_text += "[단락]\n"
    summary_text += "-" * 60 + "\n"
    if summary.paragraphs and len(summary.paragraphs) > 0:
        for para in summary.paragraphs:
            start_time = format_time(para.start)
            end_time = format_time(para.end)
            summary_text += f"\n<{para.subject}> [{start_time} ~ {end_time}]\n"
            summary_text += f"{para.summary}\n"
    else:
        summary_text += "(단락 없음)\n"
    
    summary_text += "\n" + "=" * 60 + "\n"
    
    return summary_text
